"""
Export Manager for Real-time STT
Manager eksportu dla Real-time STT

Autor: AI Assistant
Data: 2025-01-18
"""

import json
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging

from realtime_pipeline import SpeechSegment

logger = logging.getLogger(__name__)

class ExportManager:
    """
    Manager eksportu transkrypcji do różnych formatów
    """
    
    def __init__(self):
        """Inicjalizacja export managera"""
        self.supported_formats = [
            "txt", "json", "csv", "srt", "vtt", "xml", "docx"
        ]
        
    def export_transcription(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        format_type: Optional[str] = None,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ) -> bool:
        """
        Eksportuj transkrypcję do pliku
        
        Args:
            segments: Lista segmentów mowy
            output_path: Ścieżka wyjściowa
            format_type: Format eksportu (auto-detect jeśli None)
            include_metadata: Czy załączyć metadane
            include_timestamps: Czy załączyć znaczniki czasu
            include_confidence: Czy załączyć wskaźniki pewności
            
        Returns:
            True jeśli eksport się powiódł
        """
        try:
            # Auto-detect format
            if format_type is None:
                format_type = Path(output_path).suffix.lower()[1:]
            
            if format_type not in self.supported_formats:
                raise ValueError(f"Nieobsługiwany format: {format_type}")
            
            # Wybierz metodę eksportu
            export_methods = {
                "txt": self._export_txt,
                "json": self._export_json,
                "csv": self._export_csv,
                "srt": self._export_srt,
                "vtt": self._export_vtt,
                "xml": self._export_xml,
                "docx": self._export_docx
            }
            
            export_method = export_methods[format_type]
            
            # Wykonaj eksport
            export_method(
                segments, output_path,
                include_metadata=include_metadata,
                include_timestamps=include_timestamps,
                include_confidence=include_confidence
            )
            
            logger.info(f"✅ Eksport zakończony: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Błąd eksportu: {e}")
            return False
    
    def _export_txt(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do pliku tekstowego"""
        with open(output_path, 'w', encoding='utf-8') as f:
            if include_metadata:
                f.write("Real-time Speech-to-Text - Transkrypcja\n")
                f.write("=" * 50 + "\n")
                f.write(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Segmentów: {len(segments)}\n")
                f.write(f"Całkowity czas: {self._calculate_total_duration(segments):.2f}s\n")
                f.write("\n")
            
            for i, segment in enumerate(segments, 1):
                if not segment.text:
                    continue
                
                line_parts = []
                
                # Numer segmentu
                line_parts.append(f"[{i:03d}]")
                
                # Timestamp
                if include_timestamps:
                    timestamp = self._format_timestamp(segment.start_time)
                    line_parts.append(f"[{timestamp}]")
                
                # Tekst
                line_parts.append(segment.text)
                
                # Confidence
                if include_confidence and segment.transcription:
                    confidence = segment.transcription.confidence
                    line_parts.append(f"(pewność: {confidence:.2f})")
                
                f.write(" ".join(line_parts) + "\n")
                
                # Dodatkowe informacje w nowej linii
                if include_metadata and segment.transcription:
                    f.write(f"    Czas przetwarzania: {segment.transcription.processing_time:.2f}s\n")
                    f.write(f"    Długość segmentu: {segment.duration:.2f}s\n")
                
                f.write("\n")
    
    def _export_json(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do pliku JSON"""
        data = {}
        
        if include_metadata:
            data["metadata"] = {
                "export_timestamp": datetime.now().isoformat(),
                "total_segments": len(segments),
                "total_duration": self._calculate_total_duration(segments),
                "export_format": "json",
                "version": "1.0"
            }
        
        data["transcription"] = []
        
        for i, segment in enumerate(segments):
            if not segment.text:
                continue
            
            segment_data = {
                "id": i + 1,
                "text": segment.text,
                "duration": segment.duration
            }
            
            if include_timestamps:
                segment_data.update({
                    "start_time": segment.start_time,
                    "end_time": segment.end_time,
                    "timestamp_formatted": self._format_timestamp(segment.start_time)
                })
            
            if include_confidence and segment.transcription:
                segment_data.update({
                    "confidence": segment.transcription.confidence,
                    "processing_time": segment.transcription.processing_time,
                    "language": segment.transcription.language,
                    "model_used": segment.transcription.model_used
                })
            
            data["transcription"].append(segment_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def _export_csv(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do pliku CSV"""
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            # Określ kolumny
            columns = ["ID", "Text"]
            
            if include_timestamps:
                columns.extend(["Start_Time", "End_Time", "Duration", "Timestamp"])
            
            if include_confidence:
                columns.extend(["Confidence", "Processing_Time", "Language"])
            
            writer = csv.DictWriter(f, fieldnames=columns)
            writer.writeheader()
            
            for i, segment in enumerate(segments):
                if not segment.text:
                    continue
                
                row = {
                    "ID": i + 1,
                    "Text": segment.text
                }
                
                if include_timestamps:
                    row.update({
                        "Start_Time": segment.start_time,
                        "End_Time": segment.end_time,
                        "Duration": segment.duration,
                        "Timestamp": self._format_timestamp(segment.start_time)
                    })
                
                if include_confidence and segment.transcription:
                    row.update({
                        "Confidence": segment.transcription.confidence,
                        "Processing_Time": segment.transcription.processing_time,
                        "Language": segment.transcription.language
                    })
                
                writer.writerow(row)
    
    def _export_srt(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do formatu SRT (SubRip)"""
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, segment in enumerate(segments, 1):
                if not segment.text:
                    continue
                
                # Numer sekwencji
                f.write(f"{i}\n")
                
                # Znaczniki czasu SRT
                start_srt = self._format_srt_timestamp(segment.start_time)
                end_srt = self._format_srt_timestamp(segment.end_time)
                f.write(f"{start_srt} --> {end_srt}\n")
                
                # Tekst
                text = segment.text
                if include_confidence and segment.transcription:
                    confidence = segment.transcription.confidence
                    if confidence < 0.8:  # Oznacz niepewne transkrypcje
                        text = f"[?] {text}"
                
                f.write(f"{text}\n\n")
    
    def _export_vtt(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do formatu WebVTT"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("WEBVTT\n\n")
            
            if include_metadata:
                f.write("NOTE\n")
                f.write("Generated by Real-time Speech-to-Text Polish\n")
                f.write(f"Export date: {datetime.now().isoformat()}\n\n")
            
            for i, segment in enumerate(segments, 1):
                if not segment.text:
                    continue
                
                # Znaczniki czasu VTT
                start_vtt = self._format_vtt_timestamp(segment.start_time)
                end_vtt = self._format_vtt_timestamp(segment.end_time)
                f.write(f"{start_vtt} --> {end_vtt}\n")
                
                # Tekst z ewentualnymi stylami
                text = segment.text
                if include_confidence and segment.transcription:
                    confidence = segment.transcription.confidence
                    if confidence < 0.7:
                        text = f"<c.low-confidence>{text}</c>"
                    elif confidence > 0.95:
                        text = f"<c.high-confidence>{text}</c>"
                
                f.write(f"{text}\n\n")
    
    def _export_xml(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do formatu XML"""
        root = ET.Element("transcription")
        
        if include_metadata:
            metadata = ET.SubElement(root, "metadata")
            ET.SubElement(metadata, "export_date").text = datetime.now().isoformat()
            ET.SubElement(metadata, "total_segments").text = str(len(segments))
            ET.SubElement(metadata, "total_duration").text = str(self._calculate_total_duration(segments))
        
        segments_elem = ET.SubElement(root, "segments")
        
        for i, segment in enumerate(segments, 1):
            if not segment.text:
                continue
            
            segment_elem = ET.SubElement(segments_elem, "segment", id=str(i))
            
            if include_timestamps:
                ET.SubElement(segment_elem, "start_time").text = str(segment.start_time)
                ET.SubElement(segment_elem, "end_time").text = str(segment.end_time)
                ET.SubElement(segment_elem, "duration").text = str(segment.duration)
            
            ET.SubElement(segment_elem, "text").text = segment.text
            
            if include_confidence and segment.transcription:
                transcription_elem = ET.SubElement(segment_elem, "transcription_info")
                ET.SubElement(transcription_elem, "confidence").text = str(segment.transcription.confidence)
                ET.SubElement(transcription_elem, "processing_time").text = str(segment.transcription.processing_time)
                ET.SubElement(transcription_elem, "language").text = segment.transcription.language
        
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
    
    def _export_docx(
        self,
        segments: List[SpeechSegment],
        output_path: str,
        include_metadata: bool = True,
        include_timestamps: bool = True,
        include_confidence: bool = True
    ):
        """Eksport do formatu DOCX (Microsoft Word)"""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            logger.error("❌ python-docx nie zainstalowany. Użyj: pip install python-docx")
            raise ImportError("python-docx package required for DOCX export")
        
        doc = Document()
        
        # Nagłówek
        if include_metadata:
            doc.add_heading('Real-time Speech-to-Text - Transkrypcja', 0)
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"Data eksportu: ").bold = True
            info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            info_para.add_run('\n')
            
            info_para.add_run(f"Liczba segmentów: ").bold = True
            info_para.add_run(str(len(segments)))
            info_para.add_run('\n')
            
            info_para.add_run(f"Całkowity czas: ").bold = True
            info_para.add_run(f"{self._calculate_total_duration(segments):.2f}s")
            
            doc.add_page_break()
        
        # Transkrypcja
        doc.add_heading('Transkrypcja', level=1)
        
        for i, segment in enumerate(segments, 1):
            if not segment.text:
                continue
            
            # Paragraph z numerem i timestampem
            header_para = doc.add_paragraph()
            header_para.add_run(f"Segment {i}").bold = True
            
            if include_timestamps:
                timestamp = self._format_timestamp(segment.start_time)
                header_para.add_run(f" [{timestamp}]")
            
            if include_confidence and segment.transcription:
                confidence = segment.transcription.confidence
                header_para.add_run(f" (pewność: {confidence:.2f})")
            
            # Tekst transkrypcji
            text_para = doc.add_paragraph(segment.text)
            text_para.style = 'Quote'
            
            # Dodatkowe informacje
            if include_metadata and segment.transcription:
                details_para = doc.add_paragraph()
                details_para.add_run("Czas przetwarzania: ").italic = True
                details_para.add_run(f"{segment.transcription.processing_time:.2f}s")
                details_para.add_run(" | ")
                details_para.add_run("Długość: ").italic = True
                details_para.add_run(f"{segment.duration:.2f}s")
            
            doc.add_paragraph()  # Odstęp
        
        doc.save(output_path)
    
    def _calculate_total_duration(self, segments: List[SpeechSegment]) -> float:
        """Oblicz całkowity czas transkrypcji"""
        if not segments:
            return 0.0
        
        return max(segment.end_time for segment in segments) - min(segment.start_time for segment in segments)
    
    def _format_timestamp(self, timestamp: float) -> str:
        """Formatuj timestamp jako HH:MM:SS"""
        hours = int(timestamp // 3600)
        minutes = int((timestamp % 3600) // 60)
        seconds = int(timestamp % 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    
    def _format_srt_timestamp(self, timestamp: float) -> str:
        """Formatuj timestamp dla SRT (HH:MM:SS,mmm)"""
        hours = int(timestamp // 3600)
        minutes = int((timestamp % 3600) // 60)
        seconds = int(timestamp % 60)
        milliseconds = int((timestamp % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"
    
    def _format_vtt_timestamp(self, timestamp: float) -> str:
        """Formatuj timestamp dla VTT (HH:MM:SS.mmm)"""
        hours = int(timestamp // 3600)
        minutes = int((timestamp % 3600) // 60)
        seconds = int(timestamp % 60)
        milliseconds = int((timestamp % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{milliseconds:03d}"
    
    def get_supported_formats(self) -> List[str]:
        """Pobierz listę obsługiwanych formatów"""
        return self.supported_formats.copy()
    
    def batch_export(
        self,
        segments: List[SpeechSegment],
        output_dir: str,
        formats: List[str],
        base_filename: str = "transcription"
    ) -> Dict[str, bool]:
        """
        Eksportuj do wielu formatów jednocześnie
        
        Args:
            segments: Segmenty do eksportu
            output_dir: Folder wyjściowy
            formats: Lista formatów
            base_filename: Podstawowa nazwa pliku
            
        Returns:
            Słownik z wynikami dla każdego formatu
        """
        results = {}
        output_path_obj = Path(output_dir)
        output_path_obj.mkdir(parents=True, exist_ok=True)
        
        for format_type in formats:
            output_file = output_path_obj / f"{base_filename}.{format_type}"
            
            try:
                success = self.export_transcription(
                    segments, str(output_file), format_type
                )
                results[format_type] = success
                
            except Exception as e:
                logger.error(f"❌ Batch export error for {format_type}: {e}")
                results[format_type] = False
        
        return results


# Singleton instance
export_manager = ExportManager()
